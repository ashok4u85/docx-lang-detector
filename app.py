import os, re
from flask import Flask, request, jsonify, send_file, render_template
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import RGBColor, Pt
from langdetect import detect_langs, DetectorFactory
from langdetect.lang_detect_exception import LangDetectException
import tempfile

DetectorFactory.seed = 0

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 20 * 1024 * 1024

LANG_NAMES = {
    'de': 'German', 'fr': 'French', 'es': 'Spanish', 'it': 'Italian',
    'pt': 'Portuguese', 'nl': 'Dutch', 'ja': 'Japanese', 'zh-cn': 'Chinese',
    'zh-tw': 'Chinese', 'ar': 'Arabic', 'ru': 'Russian', 'ko': 'Korean',
    'pl': 'Polish', 'sv': 'Swedish', 'da': 'Danish', 'no': 'Norwegian',
    'fi': 'Finnish', 'cs': 'Czech', 'ro': 'Romanian', 'hu': 'Hungarian',
    'tr': 'Turkish', 'el': 'Greek', 'he': 'Hebrew', 'uk': 'Ukrainian',
    'hi': 'Hindi', 'id': 'Indonesian', 'th': 'Thai', 'vi': 'Vietnamese',
}

SKIP = [
    r'^\d[\d\-T:\+\./ ]*$',
    r'^[A-Z]{1,8}[\d\-_]+$',
    r'^\([\w ]+\)$',
    r'^[A-Z0-9\-\._/]{3,}$',
    r'^\d+\.\d+$',
    r'^(NI|UNK|true|false|ichicsr|Yes|No)$',
    r'^(N\.|C\.|D\.|E\.|G\.|H\.)',
    r'^\d{4}-\d{2}-\d{2}',
    r'^Page \d',
    r'^[A-Z][a-z]+(\s[A-Z]?[a-z]+){0,3}$',
]

ENGLISH_WHITELIST = {
    'element name', 'element number', 'data', 'sender type', 'sender organization',
    'patient age', 'patient weight', 'patient height', 'patient gender',
    'batch number', 'reporter country', 'reporter qualification', 'primary source',
    'case category', 'local receipt date', 'latest local receipt date',
    'worldwide unique case number', 'primary source type', 'results in death',
    'life threatening', 'disabling / incapacitating', 'congenital anomaly / birth defect',
    'other medically important condition', 'term highlighted by the reporter',
    'date of start of reaction / event', 'date of end of reaction / event',
    'outcome of reaction / event at the time of last observation',
    'identification of the country where the reaction / event occurred',
    'additional documents available?', 'documents held by sender',
    'medicinal product name as reported by the primary source',
    'substance / specified substance name',
    'batch / lot number', 'characterisation of drug role',
    'mpid version date / number', 'phpid version date/number',
    'medicinal product identifier (mpid)', 'pharmaceutical product identifier (phpid)',
    'substance/specified substance termid version date/number',
    'meddra version for reaction / event', 'reaction / event (meddra code)',
    'reaction / event as reported by the primary source in native language',
    'reaction / event as reported by the primary source language',
    'meddra version for indication', 'indication (meddra code)',
    "sender's comments", 'adult (5)', 'female (2)', 'years (a)',
    'regulatory authority (2)', 'regulator (1)', 'primary (1)',
    'consumer or other non health professional (5)',
    'ill feeling (10021345)', 'injection site pain (10022086)',
    'injection site redness (10022098)', 'injection site induration (10022075)',
    'prophylactic vaccination (10036897)', 'recovering/resolving (2)',
    'recovered/resolved (1)', 'spontaneous report(1)',
    'no information (ni)', 'unknown (unk)',
    'source of this case safety report in e2b(r2) format',
    'paul-ehrlich-institut', 'germany (de)', 'hepatitis a-virusprotein', 'hepatitis a-impfstoff',
}

def should_skip(text):
    t = text.strip()
    if len(t) < 8:
        return True
    if t.lower() in ENGLISH_WHITELIST:
        return True
    if any(re.match(p, t, re.IGNORECASE) for p in SKIP):
        return True
    alpha = [c for c in t if c.isalpha()]
    if alpha and sum(1 for c in alpha if c.isupper()) / len(alpha) > 0.5:
        return True
    return False

def detect(text):
    if should_skip(text):
        return None
    try:
        langs = detect_langs(text.strip())
        top = langs[0]
        if top.lang == 'en' or top.prob < 0.85:
            return None
        if len(text) < 20 and all(ord(c) < 128 for c in text):
            return None
        lang_name = LANG_NAMES.get(top.lang, top.lang.upper())
        return {'language': lang_name, 'confidence': round(top.prob * 100, 1)}
    except LangDetectException:
        return None

def highlight_run(run):
    rpr = run._r.get_or_add_rPr()
    hl = OxmlElement('w:highlight')
    hl.set(qn('w:val'), 'yellow')
    rpr.append(hl)

def add_lang_tag(para, lang, conf):
    run = para.add_run(f'  [\u26a0 {lang} | {conf}%]')
    run.bold = True
    run.font.color.rgb = RGBColor(0xE6, 0x00, 0x54)
    run.font.size = Pt(9)

def process_docx(filepath):
    doc = Document(filepath)
    findings = []

    def process_para(para):
        text = para.text.strip()
        if not text:
            return
        result = detect(text)
        if result:
            for run in para.runs:
                highlight_run(run)
            add_lang_tag(para, result['language'], result['confidence'])
            findings.append({'text': text[:80], 'language': result['language'], 'confidence': result['confidence']})

    for para in doc.paragraphs:
        process_para(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    process_para(para)

    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run('Non-English Language Detection Report')
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x3E, 0x01, 0x6F)

    if not findings:
        doc.add_paragraph('No non-English content detected.')
    else:
        doc.add_paragraph(f'Total instances found: {len(findings)}').runs[0].bold = True
        for i, f in enumerate(findings, 1):
            doc.add_paragraph(f"{i}. [{f['language']} | {f['confidence']}%]  {f['text']}")

    out = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(out.name)
    return out.name, findings

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/scan', methods=['POST'])
def scan():
    if 'file' not in request.files:
        return jsonify({'error': 'No file uploaded'}), 400
    f = request.files['file']
    if not f.filename.endswith('.docx'):
        return jsonify({'error': 'Only .docx files are supported'}), 400
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    f.save(tmp.name)
    try:
        out_path, findings = process_docx(tmp.name)
        os.unlink(tmp.name)
        return jsonify({'findings': findings, 'output_file': out_path})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/download', methods=['POST'])
def download():
    if 'file' not in request.files:
        return jsonify({'error': 'No file'}), 400
    f = request.files['file']
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    f.save(tmp.name)
    out_path, _ = process_docx(tmp.name)
    os.unlink(tmp.name)
    original_name = f.filename.replace('.docx', '')
    return send_file(out_path, as_attachment=True,
                     download_name=f'{original_name}_LangDetected.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)
